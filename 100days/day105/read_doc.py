from docx import Document

doc1 = Document('word1.docx')

# 读取每段内容
plist = [paragraph.text for paragraph in doc1.paragraphs]

# 遍历所有段落
for p in plist:
    print(p)

print('-----------------------------------------')


doc2 = Document('word2.docx')

plist2 = [p.text for p in doc2.paragraphs]
for p in plist2:
    print(p)


# 读取表格数据
tables = [t for t in doc2.tables]

for tab in tables:
    for row in tab.rows:
        for cell in row.cells:
            print(cell.text, end = ' ')
        print()
    print('\n')
